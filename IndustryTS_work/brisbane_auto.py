import pyautogui, sys, time
import tkinter as tk
from tkinter import messagebox
import docx
from docx.shared import Pt
from docx2pdf import convert
import os

pyautogui.FAILSAFE = True
students = []


class Student:
    def __init__(self, info, student_no):
        self.info = info
        self.student_no = student_no

    def course_sales_entry(self):
        pyautogui.click(x=-1000, y=295) 
        time.sleep(1)
        pyautogui.click(x=-996, y=260)
        time.sleep(1)
        pyautogui.moveTo(x=-1001, y=313)
        pyautogui.scroll(-900)
        pyautogui.click(x=-1644, y=275)
        
        for i in self.rsa_list:
            pyautogui.write(i)
            pyautogui.press('tab')
        pyautogui.write('4000')
        time.sleep(1)
        pyautogui.press('down')
        pyautogui.press('enter')
        pyautogui.press('tab', presses=3)
        pyautogui.write(self.dob)
        pyautogui.press('tab')
        if len(self.usi.strip()) != 10 or 'I' in self.usi or '1' in self.usi or 'O' in self.usi or '0' in self.usi:
            messagebox.showerror("USI Error", f"There is an error with {self.first_name} {self.last_name}'s USI")
            quit()
        else:
            pyautogui.write(self.usi)
        pyautogui.scroll(-900)
        pyautogui.click(x=-1434, y=255)
        time.sleep(0.5)
        pyautogui.click(x=-1560, y=415)
        time.sleep(0.1)
        pyautogui.click(x=-1607, y=392)
        if self.gender == 'Male':
            pyautogui.click(x=-1273, y=392)
        elif self.gender == 'Female':
            pyautogui.click(x=-1273, y=412)
        else:
            pyautogui.click(x=-1273, y=439)
        pyautogui.click(x=-1001, y=259)
        time.sleep(2.5)
        
class BrisStudent(Student):
    def __init__(self, info, student_no):
        self.student_no = student_no
        self.first_name = info[0]
        self.last_name = info[1]
##        info[2] = info[2].split('-')
##        self.dob = f'{info[2][0]} {info[2][1]} {info[2][2]}'
        self.dob = info[2]
        self.gender = info[3]
        self.postcode = info[10]
        self.suburb = info[9]
        if not info[6]:
            self.street_no = info[7]
        else:
            self.street_no = f'{info[6]}/{info[7]}'
        self.street = info[8]
        self.address_street = f'{self.street_no} {self.street.title()}'
        self.address_other = f'{self.suburb} QLD {self.postcode}'
        self.address = f'{self.address_street} {self.address_other}'
        self.phone = info[21]
        self.email = info[22]
        self.course = info[23]
        if "SITXFSA001" in self.course:
            self.hygiene = 1
        else:
            self.hygiene = 0
        if "SITHFAB002" in self.course:
            self.rsa = 1
        else:
            self.rsa = 0
        if "SITHGAM001" in self.course:
            self.gambling = 1
        else:
            self.gambling = 0
        self.usi = info[24]
        self.rsa_list = [self.first_name,self.last_name,self.phone
                         ,self.email,'',self.address]

    def b4_envelope(self):
        pyautogui.click(x=-1100, y=636,clicks=3)
        pyautogui.write(f'{self.first_name[0].upper()} {self.last_name}')
        pyautogui.click(x=-1100, y=669,clicks=3)
        pyautogui.write(self.address_street)
        pyautogui.click(x=-1100, y=709,clicks=3)
        pyautogui.write(self.address_other)
        pyautogui.hotkey('ctrl', 'p')
        time.sleep(2.5)
        pyautogui.press('enter')
        time.sleep(2.5)

    def certificate_creation(self):
        word = f'{self.first_name[0].upper()} {self.last_name}.docx'
##        if not self.course_number:
##            self.change_course()
        cert = docx.Document('ITS_student_certificate.docx')
        for paragraph in cert.paragraphs:
            if '{ Name }' in paragraph.text:
                paragraph.text = ''
                para = paragraph.add_run(f'{self.first_name.title()} {self.last_name.title()} \n___________________________________________')
                para.font.size = Pt(18)
                para.bold = True

            if (not self.hygiene) and ('SITXFSA001' in paragraph.text):
                paragraph.text = ''
            if (not self.rsa) and ('SITHFAB002' in paragraph.text):
                paragraph.text = ''
            if (not self.gambling) and ('SITHGAM001' in paragraph.text):
                paragraph.text = ''
                
            if '{ Date }' in paragraph.text:
                paragraph.text = paragraph.text.replace('{ Date }', '23rd of February 2022') #ask for the date
                
            if '{ course_number }' in paragraph.text:
                paragraph.text = paragraph.text.replace('{ course_number }', '51702') #ask for the course code
                
            if '{ student_number }' in paragraph.text:
                paragraph.text = paragraph.text.replace('{ student_number }', f'{self.student_no}') #ask for the first student's number
                
        cert.save(word)
        convert(word)
        os.remove(word)


class MelbStudent(Student):
    def __init__(self, info, student_no):
        self.student_no = student_no
        self.info = info
        self.first_name, self.last_name = info[1].split(" ",1)
        self.gender = info[2]
        self.dob = info[3]
        self.address = info[4]
        self.email = info[5]
        self.phone = info[6]
        self.usi = info[7]
        self.rsa_list = [self.first_name,self.last_name,self.phone
                         ,self.email,'',self.address]
        
class RMLVStudent():
    def __init__(self, info, student_no):
        self.student_no = student_no
        self._info = info
        self._RMLVlist = ['Given Names ','Surname ','Email ','DOB ','Address ','Suburb ','Postcode ']
        self._RMLVlist_variables = ['first_name','last_name','email','dob','address','suburb', 'postcode']
        for count, label in enumerate(self._RMLVlist):
            i = self._info.index(label)
            i += 1
            print(self._info[i].replace('/t',''))
            setattr(self, self._RMLVlist_variables[count], self._info[i].replace('/t',''))
            if label == 'Email ':
                self.phone = self._info[i+1]
                self.email = self.email.strip('Phone ')
        self.address_other = f'{self.suburb[2:]}QLD {self.postcode[2:]}'
                
    def rmlv_entry(self):
        pyautogui.click(x=-1000, y=295) 
        time.sleep(1)
        pyautogui.click(x=-996, y=260)
        time.sleep(1)
        pyautogui.click(x=-1698, y=934)
        time.sleep(0.1)
        pyautogui.write(self.first_name)
        pyautogui.write(self.last_name)
        pyautogui.write(self.phone)
        pyautogui.write(self.email)
        pyautogui.press('tab')
        pyautogui.write('4000')
        time.sleep(1)
        pyautogui.press('down')
        pyautogui.press('enter')
        pyautogui.press('tab', presses=2)
        pyautogui.write(self.dob)
        pyautogui.click(x=-1001, y=259)
        time.sleep(2.5)

    def b4_envelope(self):
        pyautogui.click(x=-1100, y=636,clicks=3)
        pyautogui.write(f'{self.first_name[2].upper()} {self.last_name[2:]}')
        pyautogui.click(x=-1100, y=669,clicks=3)
        pyautogui.write(self.address[2:])
        pyautogui.click(x=-1100, y=709,clicks=3)
        pyautogui.write(self.address_other)
        pyautogui.hotkey('ctrl', 'p')
        time.sleep(2.5)
        pyautogui.press('enter')
        time.sleep(2.5)

class Startup:
    def __init__(self,master):
        self.course_number = 0
        self._number = 30779
        self.master = master
        master.title('ITS Automation')
        main_frame = tk.Frame(master, bg='green', width=60,height=20)
        main_frame.pack(side=tk.TOP, fill=tk.X)
        students_button = tk.Button(main_frame, text='Enter Students', bg='SteelBlue1', fg='black', command=self.stu, width=30,height=10)
        students_button.pack()
        rsg_button = tk.Button(main_frame, text='RSA/RSG/Hygiene', bg='black', fg='white', command=self.rsa, width=30,height=10)
        rsg_button.pack()
        rmlv_button = tk.Button(main_frame, text='RMLV', bg='gold', fg='black', command=self.rmlv, width=30,height=10)
        rmlv_button.pack()
        print_envelopes = tk.Button(main_frame, text='Print Envelopes', bg='DarkOliveGreen1', fg='black', command=self.envelope, width=30,height=10)
        print_envelopes.pack()
        rsa_certs = tk.Button(main_frame, text='Create Certificates', bg='tomato', fg='black', command=self.certs, width=30,height=10)
        rsa_certs.pack()

    def certs(self):
        for student in students:
            student.certificate_creation()
        messagebox.showinfo("Done","All certificates have been created")

    def rsa(self):
        for student in students:
            student.course_sales_entry()
        messagebox.showinfo("Done","All students have been processed")

    def rmlv(self):
        for student in students:
            student.rmlv_entry()
        messagebox.showinfo("Done","All students have been processed")

    def stu(self):
        self._new = tk.Toplevel(self.master)
        new_frame = tk.Frame(self._new,width=60,height=20)
        new_frame.pack(fill=tk.BOTH)
        self.text = tk.Text(new_frame,width=60,height=10)
        self.text.pack(side=tk.TOP)
        enter_button = tk.Button(new_frame, text='Enter Brisbane', bg='black', fg='white',command=self.add_bris,width=11,height=2)
        enter_button.pack(side=tk.LEFT)
        melb_button = tk.Button(new_frame, text='Enter Melbourne', bg='gold', fg='black',command=self.add_melb,width=15,height=2)
        melb_button.pack(side=tk.LEFT)
        rmlv_button = tk.Button(new_frame, text='Enter RMLV', bg='DeepSkyBlue2', fg='black',command=self.add_rmlv,width=15,height=2)
        rmlv_button.pack(side=tk.LEFT)
        print_button = tk.Button(new_frame, text='Show Names', bg='black', fg='white', command=self.add1,width=10,height=2)
        print_button.pack(side=tk.RIGHT)
        done_button = tk.Button(new_frame, text='Done', bg='orchid1', fg='black', command=self.done,width=10,height=2)
        done_button.pack(side=tk.RIGHT)
        clear_button = tk.Button(new_frame, text='Clear', bg='navajo white', fg='black', command=self.clear,width=10,height=2)
        clear_button.pack(side=tk.RIGHT)

    def add_bris(self):
        pupils = self.text.get("1.0", tk.END).strip().splitlines()
        for pupil in pupils:
            pupil = pupil.split("\t")
            print(pupil)
            students.append(BrisStudent(pupil, self._number))
            self._number += 1
        self.text.delete("1.0", tk.END)
        

    def add_melb(self):
        pupils = self.text.get("1.0", tk.END).strip().splitlines()
        for pupil in pupils:
            pupil = pupil.split("\t")
            students.append(MelbStudent(pupil, self._number))
            self._number += 1
        self.text.delete("1.0", tk.END)
        
    def add_rmlv(self):
        pupils = self.text.get("1.0", tk.END).split('"')
        pupils = pupils[1::2]
        for pupil in pupils:
            pupil = pupil.split("\n")
            print(pupil)
            students.append(RMLVStudent(pupil, self._number))
            self._number += 1
        self.text.delete("1.0", tk.END)

    def add1(self):
        for student in students:
            print(student.first_name + str(student.student_no))

    def done(self):
        self._new.destroy()

    def clear(self):
        students.clear()
        self._number = 1
        self.text.delete("1.0", tk.END)

    def envelope(self):
        for student in students:
            student.b4_envelope()

##    def change_course(self):
##        self._course = tk.Toplevel(self.master)
##        course_frame = tk.Frame(self._course,width=60,height=20)
##        course_frame.pack(fill=tk.BOTH)
##        self.text = tk.Text(course_frame,width=60,height=10)
##        self.text.pack(side=tk.TOP)
##        course_button = tk.Button(course_frame, text='Enter Course Code', bg='black', fg='white',command=self.set_course,width=11,height=2)
##        course_button.pack(side=tk.LEFT)
##
##    def set_course(self):
##        course = self.text.get("1.0", tk.END)
##        self.course_number = course
##        self.text.delete("1.0", tk.END)
        
        

def main():
    root = tk.Tk()
    app = Startup(root)
    root.mainloop()


if __name__ == "__main__":
    main()
        
        
